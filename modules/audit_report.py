"""
Audit Report Generator
=======================
Generates a professional DOCX "Technical Audit Report" containing:
- Project metadata & cover information
- Statistical comparison tables
- All generated plots (distributions, correlations, ML comparison, feature importance)
- Privacy evaluation scorecard
- ML utility evaluation results
"""

import os
import glob
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ──────────────────────────────────────────────────────────────────────────────
# Constants
# ──────────────────────────────────────────────────────────────────────────────
PROJECT_NAME = "CTGAN Studio — Synthetic Data Generator"

REPORT_DIR = os.path.join("data", "reports")
PLOT_DIR   = os.path.join("static", "plots")


# ──────────────────────────────────────────────────────────────────────────────
# Public API
# ──────────────────────────────────────────────────────────────────────────────
def generate_audit_report(
    stats: dict,
    ml_results: dict,
    privacy: dict,
    model_type: str = "CTGAN",
    num_rows: int = 0,
    dataset_name: str = "Unknown",
    scientific_scores: dict = None,
) -> str:
    """
    Build a comprehensive Technical Audit Report and save as DOCX.

    Returns:
        Absolute path to the generated .docx file.
    """
    os.makedirs(REPORT_DIR, exist_ok=True)
    doc = Document()

    # ── Styles ────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Cover / Header ────────────────────────────────────────────────────
    _add_cover(doc, model_type, num_rows, dataset_name)

    # ── Section 1: Privacy & Security Audit ───────────────────────────────
    _add_privacy_section(doc, privacy)

    # ── Section 2: Statistical Comparison ─────────────────────────────────
    _add_stats_section(doc, stats)

    # ── Section 3: ML Evaluation ──────────────────────────────────────────
    _add_ml_section(doc, ml_results)

    # ── Section 4: Scientific Quality Scores ──────────────────────────────
    _add_scientific_section(doc, scientific_scores)

    # ── Section 5: Plots & Visualizations ─────────────────────────────────
    _add_plots_section(doc)

    # ── Footer ────────────────────────────────────────────────────────────
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Audit Report —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Save
    path = os.path.join(REPORT_DIR, "technical_audit_report.docx")
    doc.save(path)
    print(f"[Audit] Report saved → {path}")
    return os.path.abspath(path)


# ──────────────────────────────────────────────────────────────────────────────
# Private Helpers
# ──────────────────────────────────────────────────────────────────────────────
def _add_cover(doc, model_type, num_rows, dataset_name):
    """Add cover page / header section."""
    # Title
    title = doc.add_heading("Technical Audit Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("Synthetic Data Quality Assessment", level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    # Info table
    info = [
        ("Tool",               PROJECT_NAME),
        ("Date Generated",     datetime.now().strftime("%B %d, %Y — %H:%M")),
        ("Dataset",            dataset_name),
        ("Model Architecture", model_type),
        ("Synthetic Records",  str(num_rows)),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        # Bold label
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_page_break()


def _add_privacy_section(doc, privacy):
    """Add Privacy & Security Audit section."""
    doc.add_heading("1. Privacy & Security Audit", level=1)
    doc.add_paragraph(
        "This section evaluates the risk of identity disclosure by computing "
        "the Distance to Closest Record (DCR) between every synthetic record "
        "and its nearest real-world counterpart."
    )

    if not privacy:
        doc.add_paragraph("Privacy evaluation data not available.")
        return

    metrics = [
        ("Privacy Score", f"{privacy.get('privacy_score', 'N/A')}%"),
        ("Risk Level", privacy.get("risk_level", "N/A")),
        ("Mean DCR", str(privacy.get("mean_dcr", "N/A"))),
        ("Median DCR", str(privacy.get("median_dcr", "N/A"))),
        ("Min DCR (worst case)", str(privacy.get("min_dcr", "N/A"))),
        ("Records at Risk", f"{privacy.get('at_risk_count', 0)} / {privacy.get('total_synthetic', 0)} ({privacy.get('pct_at_risk', 0)}%)"),
    ]
    table = doc.add_table(rows=len(metrics), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(metrics):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()


def _add_stats_section(doc, stats):
    """Add Statistical Comparison section."""
    doc.add_heading("2. Statistical Comparison", level=1)
    doc.add_paragraph(
        "Side-by-side comparison of mean and standard deviation between "
        "real and synthetic feature distributions."
    )

    if not stats:
        doc.add_paragraph("Statistical comparison data not available.")
        return

    for label, key in [("Mean Comparison", "mean_comparison"), ("Standard Deviation Comparison", "std_comparison")]:
        df = stats.get(key)
        if df is None:
            continue
        doc.add_heading(label, level=2)
        _df_to_table(doc, df)
        doc.add_paragraph()


def _add_ml_section(doc, ml_results):
    """Add ML Evaluation section."""
    doc.add_heading("3. Machine Learning Utility Evaluation", level=1)
    doc.add_paragraph(
        "Each classifier is trained twice — once on the real data and once "
        "on the synthetic data — then tested on a held-out portion of the "
        "real dataset. Close performance indicates high data utility."
    )

    if not ml_results:
        doc.add_paragraph("ML evaluation data not available.")
        return

    for source, key in [("Trained on Real Data", "real_results"), ("Trained on Synthetic Data", "synthetic_results")]:
        doc.add_heading(source, level=2)
        results = ml_results.get(key, [])
        if not results:
            continue
        # Build a table: Model | Accuracy | Precision | Recall | F1 | ROC-AUC
        cols = ["Model", "Accuracy", "Precision", "Recall", "F1 Score", "ROC-AUC"]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Light Grid Accent 1"
        for i, col in enumerate(cols):
            table.rows[0].cells[i].text = col
            for run in table.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        for r in results:
            row_cells = table.add_row().cells
            for i, col in enumerate(cols):
                row_cells[i].text = str(r.get(col, ""))
        doc.add_paragraph()

    # Feature Importance
    fi = ml_results.get("feature_importance")
    if fi:
        doc.add_heading("Feature Importance Alignment (Random Forest)", level=2)
        doc.add_paragraph(
            "Compares which features the model considers most important "
            "when trained on real vs. synthetic data."
        )
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        headers = ["Feature", "Real Importance", "Synthetic Importance", "Δ Difference"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            for run in table.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        for j, feat in enumerate(fi["features"]):
            row_cells = table.add_row().cells
            row_cells[0].text = feat
            row_cells[1].text = str(fi["real"][j])
            row_cells[2].text = str(fi["synthetic"][j])
            row_cells[3].text = str(round(abs(fi["real"][j] - fi["synthetic"][j]), 4))
        doc.add_paragraph()


def _add_scientific_section(doc, scientific_scores):
    """Add Scientific Quality Scores section."""
    doc.add_heading("4. Scientific Quality Scores", level=1)
    doc.add_paragraph(
        "Formal statistical tests quantify how closely the synthetic "
        "distributions match the original data."
    )

    if not scientific_scores:
        doc.add_paragraph("Scientific quality scores not available.")
        return

    # Summary table
    summary = [
        ("KS Similarity (1.0 = identical)", str(scientific_scores.get("ks_mean_score", "N/A"))),
        ("JS Similarity (1.0 = zero divergence)", str(scientific_scores.get("js_mean_score", "N/A"))),
        ("Correlation Similarity", f"{scientific_scores.get('corr_similarity', 'N/A')}%"),
    ]
    table = doc.add_table(rows=len(summary), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(summary):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    # KS Test per column
    ks_results = scientific_scores.get("ks_results", [])
    if ks_results:
        doc.add_heading("Kolmogorov-Smirnov Test (per column)", level=2)
        cols = ["Column", "KS Statistic", "P-Value", "Result"]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Light Grid Accent 1"
        for i, c in enumerate(cols):
            table.rows[0].cells[i].text = c
            for run in table.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        for r in ks_results:
            row_cells = table.add_row().cells
            row_cells[0].text = r["column"]
            row_cells[1].text = str(r["statistic"])
            row_cells[2].text = str(r["p_value"])
            row_cells[3].text = "Pass" if r["pass"] else "Fail"

    # JS Distance per column
    js_results = scientific_scores.get("js_results", [])
    if js_results:
        doc.add_heading("Jensen-Shannon Distance (per column)", level=2)
        cols = ["Column", "JS Distance", "Similarity"]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Light Grid Accent 1"
        for i, c in enumerate(cols):
            table.rows[0].cells[i].text = c
            for run in table.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        for r in js_results:
            row_cells = table.add_row().cells
            row_cells[0].text = r["column"]
            row_cells[1].text = str(r["distance"])
            row_cells[2].text = str(round(1 - r["distance"], 4))

    doc.add_paragraph()


def _add_plots_section(doc):
    """Embed all saved plots from the plots directory."""
    doc.add_heading("5. Distribution & Correlation Plots", level=1)
    doc.add_paragraph(
        "The following plots provide visual validation of the synthetic "
        "data quality against the original distribution."
    )

    plot_files = sorted(glob.glob(os.path.join(PLOT_DIR, "*.png")))
    if not plot_files:
        doc.add_paragraph("No plot files found.")
        return

    for pf in plot_files:
        name = os.path.basename(pf).replace("_", " ").replace(".png", "").title()
        doc.add_heading(name, level=3)
        try:
            doc.add_picture(pf, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Could not embed: {e}]")
        doc.add_paragraph()


def _df_to_table(doc, df):
    """Convert a pandas DataFrame to a Word table."""
    table = doc.add_table(rows=1, cols=len(df.columns) + 1)
    table.style = "Light Grid Accent 1"

    # Header
    table.rows[0].cells[0].text = ""
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i + 1].text = str(col)
        for run in table.rows[0].cells[i + 1].paragraphs[0].runs:
            run.bold = True

    # Rows
    for idx, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(idx)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
        for i, val in enumerate(row):
            cells[i + 1].text = f"{val:.4f}" if isinstance(val, float) else str(val)
