"""
Synthetic Data Generator Module
================================
Generates synthetic records from a trained CTGAN model and
exports the result in multiple formats (CSV, Excel, JSON, DOCX).
"""

import os
import pandas as pd
from docx import Document


def generate_synthetic_data(model, num_rows: int) -> pd.DataFrame:
    """
    Sample synthetic rows from a trained CTGAN model.

    Args:
        model:    A fitted CTGANSynthesizer.
        num_rows: Number of synthetic records to generate.

    Returns:
        A pandas DataFrame of synthetic data.
    """
    print(f"[Generator] Generating {num_rows} synthetic rows …")
    synthetic_df = model.sample(num_rows=num_rows)
    print("[Generator] Generation complete.")
    return synthetic_df


def export_data(
    df: pd.DataFrame,
    fmt: str,
    output_dir: str = os.path.join("data", "synthetic"),
    filename: str = "synthetic_data",
) -> str:
    """
    Export a DataFrame to the specified format.

    Supported formats: csv, xlsx (Excel), json, docx.

    Args:
        df:         DataFrame to export.
        fmt:        One of 'csv', 'xlsx', 'json', 'docx'.
        output_dir: Directory to write the file to.
        filename:   Base filename (without extension).

    Returns:
        Absolute path of the exported file.

    Raises:
        ValueError: If an unsupported format is requested.
    """
    os.makedirs(output_dir, exist_ok=True)
    fmt = fmt.lower().strip()

    if fmt == "csv":
        path = os.path.join(output_dir, f"{filename}.csv")
        df.to_csv(path, index=False)

    elif fmt in ("xlsx", "excel"):
        path = os.path.join(output_dir, f"{filename}.xlsx")
        df.to_excel(path, index=False, engine="openpyxl")

    elif fmt == "json":
        path = os.path.join(output_dir, f"{filename}.json")
        df.to_json(path, orient="records", indent=2)

    elif fmt in ("docx", "doc"):
        path = os.path.join(output_dir, f"{filename}.docx")
        _export_docx(df, path)

    else:
        raise ValueError(f"Unsupported export format: {fmt}")

    print(f"[Generator] Exported → {path}")
    return os.path.abspath(path)


def _export_docx(df: pd.DataFrame, path: str) -> None:
    """
    Write the DataFrame into a Word document as a table.

    Args:
        df:   DataFrame to export.
        path: Output .docx file path.
    """
    doc = Document()
    doc.add_heading("Synthetic Data Report", level=1)

    # Add a summary paragraph
    doc.add_paragraph(
        f"This report contains {len(df)} synthetic records "
        f"with {len(df.columns)} columns."
    )

    # Create table (header + rows)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"

    # Header row
    for i, col_name in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col_name)

    # Data rows (limit to 100 rows for readability)
    for _, row in df.head(100).iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    if len(df) > 100:
        doc.add_paragraph(
            f"(Showing first 100 of {len(df)} rows.)"
        )

    doc.save(path)
